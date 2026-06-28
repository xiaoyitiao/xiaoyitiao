from docx import Document
from docx.shared import Pt

doc = Document()
doc.add_heading('新闻2班_小组一_实践二中期汇报', level=1)

def add_heading(text, level=2):
    doc.add_paragraph(text, style=f'Heading {level}')


def add_paragraph(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

add_heading('一、选题回顾', 2)
add_paragraph('选题名称：智药伴 - AI 老年人用药提醒 Web 应用')
add_paragraph('真实需求：帮助老年人及其家属解决每日用药提醒不及时、服药记录难管理、药物信息难查询的问题，提升用药安全性与家属关怀体验。')
add_paragraph('作品目标：构建一款面向老年人、兼顾家属监督的移动端友好用药提醒系统。')
add_paragraph('阶段进展：已进入开发阶段，完成页面结构、基础数据存储、语音录入、拍照识别、AI助手、家属关怀与设置功能的核心实现。')

add_heading('二、当前进展', 2)
add_paragraph('已完成的页面与功能：')
add_paragraph('1. 今日用药页面：展示当天待服药品列表、用药进度、打卡完成功能。')
add_paragraph('2. 添加药品页面：实现语音录入、拍照识别和手动输入三种添加方式。')
add_paragraph('3. 用药计划页面：展示所有已添加药品并支持删除操作。')
add_paragraph('4. AI用药助手：提供本地知识库问答，并支持未来接入真实大模型 API。')
add_paragraph('5. 家属关怀页面：展示今日应服、已服、漏服统计，并支持远程云端查询。')
add_paragraph('6. 设置与长辈模式：支持大字体高对比、语音播报提醒、云端同步配置与AI接口测试功能。')
add_paragraph('演示说明：当前版本可在浏览器中运行，已完成可交互的核心模块，可进行药品添加、今日打卡、AI问答、长辈模式切换。')

add_heading('三、技术方案', 2)
add_paragraph('前端技术：')
add_paragraph('1. HTML5 + CSS3 + JavaScript 原生开发，使用 Vite 构建提升开发效率。')
add_paragraph('2. 响应式移动端布局，页面适配老年人浏览体验。')
add_paragraph('3. localStorage 本地持久化保存药品列表、打卡记录与设置。')
add_paragraph('AI与辅助技术：')
add_paragraph('1. 语音识别与播报：使用 Web Speech API 实现语音录入、AI助手语音问答和提醒播报。')
add_paragraph('2. 拍照识别：使用 Tesseract.js 在浏览器内执行 OCR，提取药品包装文字并结合本地药品库智能解析。')
add_paragraph('3. AI助手：构建本地知识库问答，支持“启用真实 API”模式调用大模型接口进行高级问答与文字解析。')
add_paragraph('4. 云端同步：使用 LeanCloud 实现老人端上传数据与家属端远程查看功能。')
add_paragraph('关键功能实现：')
add_paragraph('1. 语音录入模块解析自然语言药品信息，并自动提取时间、剂量、备注等字段。')
add_paragraph('2. 拍照识别模块通过 OCR 提取药品文字，再匹配本地数据库或调用 AI 解析，完成药品信息自动填充。')
add_paragraph('3. 家属关怀页面以日常服药统计为核心，自动计算漏服提醒，并通过云同步支持远程查看。')

add_heading('四、难点与解决', 2)
add_paragraph('目前遇到的主要问题：')
add_paragraph('1. 语音识别与中文自然语言解析：老年人语音描述形式多样，解析药品名称和服药时间具有一定难度。')
add_paragraph('2. OCR 识别准确率：药品照片质量和中文字体复杂，直接提取药品名称与用法信息不稳定。')
add_paragraph('3. AI 接口接入与安全：前端调用真实 API 需避免密钥泄露，并兼顾本地降级体验。')
add_paragraph('4. 老年人友好交互：界面需兼顾大字体、高对比、简洁按钮与提示语。')
add_paragraph('解决方案与计划：')
add_paragraph('1. 采用 Web Speech API 进行语音识别，并使用规则解析与本地提示降低识别错误影响。')
add_paragraph('2. 优先使用本地药品库匹配 OCR 结果，未匹配时可调用 AI 辅助解析，保证识别失败时仍有可用输入方式。')
add_paragraph('3. 提供“启用真实 API”开关，默认使用本地知识库问答；真实接口密钥仅在设置页保存，避免直接硬编码。')
add_paragraph('4. 实现长辈模式与语音播报，让操作更适合老年人使用。')

add_heading('五、后续计划', 2)
add_paragraph('剩余工作与下一步安排：')
add_paragraph('1. 优化 UI 细节与操作流程，增强页面可访问性与交互提示。')
add_paragraph('2. 提高 OCR 识别与语音解析准确率，补齐更多药品名称词库与解析规则。')
add_paragraph('3. 补充家属端远程查看与云同步稳定性，完善数据同步错误提示。')
add_paragraph('4. 增加药品安全提醒、用药日志导出、定时通知等功能。')
add_paragraph('5. 准备课堂汇报材料（PPT/PDF），并整理源码压缩包提交。')

add_heading('作品链接或仓库地址', 2)
add_paragraph('本地仓库地址：c:\\Users\\Admin\\zhiyaoban-app')
add_paragraph('当前无线上访问链接，可通过本地开发环境运行 `npm install` 后执行 `npm run dev` 进行演示。')


doc.save('新闻2班_小组一_实践二中期汇报.docx')
print('saved')
